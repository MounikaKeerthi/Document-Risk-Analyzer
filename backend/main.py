from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
from sqlalchemy import or_
from datetime import datetime
from typing import List
from io import BytesIO
from pypdf import PdfReader
from docx import Document as DocxDocument

from database import Base, engine, SessionLocal
from models import Document
from schemas import DocumentCreate, DocumentUpdate, DocumentResponse
from analyzer import analyze_document

Base.metadata.create_all(bind=engine)

app = FastAPI(title="PrivyDoc AI API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
    "http://localhost:5173",
    "https://document-risk-analyzer-frontend.onrender.com",
],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

def extract_text_from_upload(file: UploadFile, file_bytes: bytes) -> str:
    filename = file.filename.lower()

    if filename.endswith(".txt"):
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            raise HTTPException(status_code=400, detail="Could not decode TXT file.")

    if filename.endswith(".pdf"):
        try:
            reader = PdfReader(BytesIO(file_bytes))
            text_parts = []

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            text = "\n".join(text_parts).strip()

            if not text:
                raise HTTPException(
                    status_code=400,
                    detail="Could not extract text from PDF. Scanned PDFs are not supported yet."
                )

            return text
        except Exception as error:
            raise HTTPException(status_code=400, detail=f"PDF extraction failed: {str(error)}")

    if filename.endswith(".docx"):
        try:
            doc = DocxDocument(BytesIO(file_bytes))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs]).strip()

            if not text:
                raise HTTPException(status_code=400, detail="Could not extract text from DOCX.")

            return text
        except Exception as error:
            raise HTTPException(status_code=400, detail=f"DOCX extraction failed: {str(error)}")

    raise HTTPException(
        status_code=400,
        detail="Unsupported file type. Upload .txt, .pdf, or .docx."
    )

@app.get("/")
def home():
    return {"message": "PrivyDoc AI backend is running"}

@app.post("/documents", response_model=DocumentResponse)
def create_document(document: DocumentCreate, db: Session = Depends(get_db)):
    new_document = Document(
        title=document.title,
        document_type=document.document_type,
        client_name=document.client_name,
        content=document.content
    )

    db.add(new_document)
    db.commit()
    db.refresh(new_document)

    return new_document

@app.post("/documents/upload", response_model=DocumentResponse)
async def upload_document(
    title: str = Form(...),
    document_type: str = Form(...),
    client_name: str = Form(None),
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    file_bytes = await file.read()
    content = extract_text_from_upload(file, file_bytes)

    new_document = Document(
        title=title,
        document_type=document_type,
        client_name=client_name,
        content=content
    )

    db.add(new_document)
    db.commit()
    db.refresh(new_document)

    return new_document

@app.get("/documents", response_model=List[DocumentResponse])
def get_documents(
    query: str = "",
    document_type: str = "",
    risk_level: str = "",
    db: Session = Depends(get_db)
):
    documents_query = db.query(Document)

    if query:
        search = f"%{query}%"
        documents_query = documents_query.filter(
            or_(
                Document.title.ilike(search),
                Document.client_name.ilike(search),
                Document.content.ilike(search)
            )
        )

    if document_type:
        documents_query = documents_query.filter(Document.document_type == document_type)

    if risk_level:
        documents_query = documents_query.filter(Document.risk_level == risk_level)

    return documents_query.order_by(Document.created_at.desc()).all()

@app.get("/documents/{document_id}", response_model=DocumentResponse)
def get_document(document_id: int, db: Session = Depends(get_db)):
    document = db.query(Document).filter(Document.id == document_id).first()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    return document

@app.put("/documents/{document_id}", response_model=DocumentResponse)
def update_document(
    document_id: int,
    updated_document: DocumentUpdate,
    db: Session = Depends(get_db)
):
    document = db.query(Document).filter(Document.id == document_id).first()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    document.title = updated_document.title
    document.document_type = updated_document.document_type
    document.client_name = updated_document.client_name
    document.content = updated_document.content
    document.updated_at = datetime.utcnow()

    db.commit()
    db.refresh(document)

    return document

@app.delete("/documents/{document_id}")
def delete_document(document_id: int, db: Session = Depends(get_db)):
    document = db.query(Document).filter(Document.id == document_id).first()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    db.delete(document)
    db.commit()

    return {"message": "Document deleted successfully"}

@app.post("/documents/{document_id}/analyze", response_model=DocumentResponse)
def analyze_existing_document(document_id: int, db: Session = Depends(get_db)):
    document = db.query(Document).filter(Document.id == document_id).first()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    try:
        analysis = analyze_document(document.content)
    except Exception as error:
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(error)}")

    document.risk_level = analysis["risk_level"]
    document.summary = analysis["summary"]
    document.risky_clauses = analysis["risky_clauses"]
    document.privacy_concerns = analysis["privacy_concerns"]
    document.suggested_questions = analysis["suggested_questions"]
    document.updated_at = datetime.utcnow()

    db.commit()
    db.refresh(document)

    return document
